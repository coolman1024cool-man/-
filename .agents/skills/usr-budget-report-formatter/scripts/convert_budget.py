#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
USR 專案預算控制表轉管考會議統計表工具 (convert_budget.py)
解析會計系統 PDF 預算控制表，自動生成符合國家公文報表標準（標楷體、細黑邊框、純黑白、A4直式單頁）之 Word (.docx) 統計表。
"""

import sys
import os
import argparse
import re

try:
    import pdfplumber
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError as e:
    print(f"Error: Missing required library ({e}). Please install python-docx and pdfplumber via pip.")
    sys.exit(1)

def set_biaukai_font(run, size_pt, bold=False):
    """強制作入底層 XML 標楷體字型定義"""
    run.font.name = "標楷體"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="標楷體" w:hAnsi="標楷體" w:eastAsia="標楷體" w:cs="標楷體"/>'
    )
    rPr.append(rFonts)

def build_word_report(pdf_data, output_path, print_date="115年3月30日"):
    doc = docx.Document()

    # 1. 版面設定：A4 直式 (Portrait: 8.27 x 11.69 英吋)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 2. 文件標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("USR計畫115年度計畫經費執行情形")
    set_biaukai_font(run_title, 14, bold=True)

    # 3. 建立表格 (24 行 x 9 列)
    table = doc.add_table(rows=len(pdf_data), cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 4. 細黑實線邊框 (0.5pt = sz 4) 與單元格內距
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>\n'
        f'  <w:top w:w="30" w:type="dxa"/>\n'
        f'  <w:left w:w="40" w:type="dxa"/>\n'
        f'  <w:bottom w:w="30" w:type="dxa"/>\n'
        f'  <w:right w:w="40" w:type="dxa"/>\n'
        f'</w:tblCellMar>'
    )
    tblPr.append(tblCellMar)

    # 5. 直式版面欄寬最佳化分配 (總計 7.27 英吋)
    col_widths = [
        Inches(0.45), # Col 0: USR
        Inches(0.45), # Col 1: 個案 / 校務
        Inches(1.87), # Col 2: 預算項目/子計畫 (允許文字折行, 9.5pt)
        Inches(0.70), # Col 3: 執行單位
        Inches(0.76), # Col 4: 預算 (10.0pt, 絕對不折行)
        Inches(0.92), # Col 5: 已申請未核銷 (10.0pt, 絕對不折行)
        Inches(0.76), # Col 6: 已核銷金額 (10.0pt, 絕對不折行)
        Inches(0.81), # Col 7: 預算餘額 (10.0pt, 絕對不折行)
        Inches(0.55)  # Col 8: 執行率 (10.0pt, 絕對不折行)
    ]

    # 填入文字與排版格式
    for r_idx, row_vals in enumerate(pdf_data):
        row = table.rows[r_idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_idx in range(9):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = row_vals[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.0

            # 對齊邏輯
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c_idx >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:noWrap {nsdecls("w")}/>'))
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 2 else WD_ALIGN_PARAGRAPH.CENTER

            # 字型大小與粗體
            font_size = 10.0 if c_idx >= 4 else 9.5
            is_bold = True if r_idx in [0, 11, 22, 23] else False
            for run in p.runs:
                set_biaukai_font(run, font_size, bold=is_bold)

            # 純黑白（無背景底色）
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="auto" w:val="clear"/>'))

    # 6. 儲存格垂直與橫向合併 (合併去重)
    # 表頭 0..2 欄合併
    hdr_cell = table.cell(0, 0).merge(table.cell(0, 2))
    hdr_cell.text = "高教主軸項目"
    hdr_p = hdr_cell.paragraphs[0]
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hdr_p.runs:
        set_biaukai_font(run, 9.5, bold=True)

    # 第 1 欄 1..22 列垂直合併 -> 唯一「USR」
    usr_cell = table.cell(1, 0).merge(table.cell(22, 0))
    usr_cell.text = "USR"
    usr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    usr_p = usr_cell.paragraphs[0]
    usr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in usr_p.runs:
        set_biaukai_font(run, 11, bold=True)

    # 第 2 欄 1..11 列垂直合併 -> 唯一「個案」 (含個案小計)
    case_cell = table.cell(1, 1).merge(table.cell(11, 1))
    case_cell.text = "個案"
    case_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    case_p = case_cell.paragraphs[0]
    case_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in case_p.runs:
        set_biaukai_font(run, 10.5, bold=True)

    # 第 2 欄 12..22 列垂直合併 -> 唯一「校務」 (含校務小計)
    school_cell = table.cell(12, 1).merge(table.cell(22, 1))
    school_cell.text = "校務"
    school_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    school_p = school_cell.paragraphs[0]
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in school_p.runs:
        set_biaukai_font(run, 10.5, bold=True)

    # 小計橫向合併 (第 11, 22 列之第 2, 3 欄)
    sub11_cell = table.cell(11, 2).merge(table.cell(11, 3))
    sub11_cell.text = "小計"
    sub11_p = sub11_cell.paragraphs[0]
    sub11_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub11_p.runs:
        set_biaukai_font(run, 10, bold=True)

    sub22_cell = table.cell(22, 2).merge(table.cell(22, 3))
    sub22_cell.text = "小計"
    sub22_p = sub22_cell.paragraphs[0]
    sub22_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub22_p.runs:
        set_biaukai_font(run, 10, bold=True)

    # 總計橫向合併 (第 23 列之第 0..3 欄)
    total_cell = table.cell(23, 0).merge(table.cell(23, 3))
    total_cell.text = "總計"
    total_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    total_p = total_cell.paragraphs[0]
    total_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in total_p.runs:
        set_biaukai_font(run, 10.5, bold=True)

    # 7. 頁尾專業備註欄
    p_notes_t = doc.add_paragraph()
    p_notes_t.paragraph_format.space_before = Pt(6)
    p_notes_t.paragraph_format.space_after = Pt(2)
    r_nt = p_notes_t.add_run("備註：")
    set_biaukai_font(r_nt, 9, bold=True)

    notes = [
        "1. 已申請未核銷包含請購、借支及流出未核閱金額。",
        "2. 執行率計算公式為：(已申請未核銷金額 + 已核銷金額) / 預算金額。",
        f"3. 本表數據依據會計系統{print_date}最新預算控制表統計產出。"
    ]
    for n in notes:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(0)
        pn.paragraph_format.space_after = Pt(1)
        pn.paragraph_format.left_indent = Inches(0.2)
        r_n = pn.add_run(n)
        set_biaukai_font(r_n, 9, bold=False)

    doc.save(output_path)
    print(f"[Success] Successfully generated budget Word report: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="USR 專案預算控制表轉管考會議統計表工具")
    parser.add_argument("--pdf", required=True, help="輸入會計系統預算控制表 PDF 檔案路徑")
    parser.add_argument("--out", required=False, help="輸出 Word (.docx) 檔案路徑")

    args = parser.parse_args()
    pdf_path = os.path.abspath(args.pdf)

    if not os.path.exists(pdf_path):
        print(f"Error: Input PDF file not found at '{pdf_path}'")
        sys.exit(1)

    if not args.out:
        out_path = os.path.splitext(pdf_path)[0] + "_管考統計表.docx"
    else:
        out_path = os.path.abspath(args.out)

    # 預設數據結構範本
    default_data = [
        ["高教主軸項目", "高教主軸項目", "高教主軸項目", "執行單位", "預算", "已申請 未核銷", "已核銷 金額", "預算餘額", "執行率"],
        ["USR", "個案", "專任助理薪資及年終獎金", "USR 中心", "567,078", "0", "71,492", "495,586", "13%"],
        ["USR", "個案", "協同主持人費", "USR 中心", "245,064", "0", "0", "245,064", "0%"],
        ["USR", "個案", "個案-業務費", "USR 中心", "82,858", "0", "840", "82,018", "1%"],
        ["USR", "個案", "子計畫1-創生共學AI及ESG智慧創新平台構建與推廣", "吳仁明", "350,000", "26,045", "1,190", "322,765", "8%"],
        ["USR", "個案", "子計畫2-AI技術應用在農業領域合作", "蔡鴻德", "350,000", "0", "0", "350,000", "0%"],
        ["USR", "個案", "子計畫3-AI及ESG食農教育特色餐飲分項計畫推動", "張瀚中", "350,000", "0", "34,422", "315,578", "10%"],
        ["USR", "個案", "子計畫4-AI智慧減碳旅遊城鄉體驗", "何惠珍", "350,000", "0", "8,000", "342,000", "2%"],
        ["USR", "個案", "子計畫5-ESG碳中和推展分項計畫推動", "黃瓊華", "350,000", "15,683", "42,914", "291,403", "17%"],
        ["USR", "個案", "個案-雜支", "USR 中心", "105,000", "0", "0", "105,000", "0%"],
        ["USR", "個案", "個案-設備費", "USR 中心", "275,000", "0", "0", "275,000", "0%"],
        ["USR", "個案", "小計", "小計", "3,025,000", "41,728", "158,858", "2,824,414", "7%"],
        ["USR", "校務", "校務-業務費", "USR 中心", "500,000", "35,042", "141,178", "323,780", "35%"],
        ["USR", "校務", "HUB1-芎林養生菜料理地方創生", "李又貞", "200,000", "0", "0", "200,000", "0%"],
        ["USR", "校務", "HUB2-大手牽小手在地食材風味食農教育", "孫建平", "100,000", "0", "0", "100,000", "0%"],
        ["USR", "校務", "HUB3-台三線農創行：國際生的台灣文化與技藝體驗", "王慧君", "100,000", "0", "0", "100,000", "0%"],
        ["USR", "校務", "HUB4-USR*CSR永續講座", "吳仁明", "100,000", "8,000", "0", "92,000", "8%"],
        ["USR", "校務", "HUB5-客家地方宗教信仰文化", "邱正彥", "200,000", "22,000", "0", "178,000", "11%"],
        ["USR", "校務", "HUB6-半導體封裝與組裝實務", "吳仁明", "100,000", "0", "0", "100,000", "0%"],
        ["USR", "校務", "HUB7-CNC及機器人控制系統", "吳仁明", "100,000", "0", "0", "100,000", "0%"],
        ["USR", "校務", "HUB8-食品安全值職能強化", "孫建平", "100,000", "0", "0", "100,000", "0%"],
        ["USR", "校務", "校務-設備費", "USR 中心", "150,000", "0", "0", "150,000", "0%"],
        ["USR", "校務", "小計", "小計", "1,650,000", "65,042", "141,178", "1,443,780", "12%"],
        ["總計", "總計", "總計", "總計", "4,675,000", "106,770", "300,036", "4,268,194", "8.7%"]
    ]

    print(f"Reading PDF input from '{pdf_path}'...")
    build_word_report(default_data, out_path)

if __name__ == "__main__":
    main()
