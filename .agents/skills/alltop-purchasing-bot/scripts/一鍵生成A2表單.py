import pandas as pd
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import datetime

tool_dir = r'D:\01邱正彥1130801-\01經費\ALLTOP_採購自動化工具'
excel_path = os.path.join(tool_dir, '採購自動化專用表.xlsx')
template_path = os.path.join(tool_dir, 'A2規格表_Template.docx')
output_dir = os.path.join(tool_dir, 'A2產出結果')

os.makedirs(output_dir, exist_ok=True)

print("正在讀取 Excel 資料庫...")
try:
    df = pd.read_excel(excel_path)
    df = df.fillna('')
except Exception as e:
    print(f"讀取 Excel 失敗: {e}")
    input("按 Enter 結束...")
    exit(1)

# 計算民國年日期
today = datetime.date.today()
roc_year = today.year - 1911
date_str = f"{roc_year}年{today.month}月{today.day}日"

def set_font(run, font_name='標楷體', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

count = 0
for index, row in df.iterrows():
    title = str(row['品名']).strip()
    if not title:
        continue
        
    spec = str(row['規格'])
    unit = str(row['數量']) + str(row['計量單位'])
    
    doc = docx.Document(template_path)
    
    # 替換段落中的佔位符
    for p in doc.paragraphs:
        for run in p.runs:
            if '{{TITLE}}' in run.text or '{{DATE}}' in run.text:
                run.text = run.text.replace('{{TITLE}}', title).replace('{{DATE}}', date_str)
                set_font(run)
                
    # 替換表格中的佔位符
    for table in doc.tables:
        for r in table.rows:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '{{UNIT}}' in run.text:
                            run.text = run.text.replace('{{UNIT}}', unit)
                            set_font(run)
                        if '{{SPEC}}' in run.text:
                            run.text = run.text.replace('{{SPEC}}', '')
                            lines = spec.split('\n')
                            run.add_text(lines[0])
                            for line in lines[1:]:
                                run.add_break()
                                run.add_text(line)
                            set_font(run)

    # 存檔
    safe_title = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ' or c=='_']).rstrip()
    if not safe_title:
        safe_title = f"Item_{index}"
    out_name = f"A2規格表-{safe_title}.docx"
    
    doc.save(os.path.join(output_dir, out_name))
    print(f"✅ 成功生成: {out_name}")
    count += 1

print("\n========================================")
print(f"🎉 全部完成！共生成了 {count} 份 A2 規格表。")
print(f"檔案已儲存於：{output_dir}")
print("========================================")
input("按 Enter 鍵關閉視窗...")
