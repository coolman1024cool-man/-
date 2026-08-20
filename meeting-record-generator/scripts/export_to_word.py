import sys
import subprocess
import os

try:
    import docx
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn

if len(sys.argv) < 3:
    print("Usage: python export_to_word.py <input_md> <output_docx>")
    sys.exit(1)

md_path = sys.argv[1]
out_path = sys.argv[2]

if not os.path.exists(md_path):
    print(f"Error: input file {md_path} not found.")
    sys.exit(1)

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = docx.Document()

# 1. Set Margins (Top and Bottom 2cm)
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# 2. Font configuration helper
def set_font(run, size=12, bold=False):
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    if bold:
        run.bold = True

# 3. Parse and write lines
for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('# '):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        set_font(run, size=18, bold=True)
    elif line.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line[3:])
        set_font(run, size=16, bold=True)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line[4:])
        set_font(run, size=14, bold=True)
    elif line.startswith('- [ ]'):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run('☐ ' + line[5:].strip())
        set_font(run, size=12)
    elif line.startswith('- [x]') or line.startswith('- [X]'):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run('☑ ' + line[5:].strip())
        set_font(run, size=12)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(line[2:])
        set_font(run, size=12)
    elif line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet 2')
        run = p.add_run(line[2:])
        set_font(run, size=12)
    elif line.startswith('![') and '](' in line:
        import re
        match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if match:
            img_path = match.group(1)
            if img_path.startswith('file:///'):
                img_path = img_path[8:]
            img_path = img_path.replace('%20', ' ')
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = 1 # center
                p.add_run().add_picture(img_path, width=Cm(15))
            else:
                p = doc.add_paragraph()
                p.add_run(f'[Image not found: {img_path}]')
    elif line.startswith('---'):
        p = doc.add_paragraph()
        run = p.add_run('_' * 40)
        set_font(run, size=12)
    else:
        p = doc.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            is_bold = (i % 2 == 1)
            set_font(run, size=12, bold=is_bold)

# Make sure output directory exists
os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
doc.save(out_path)
print(f'Successfully generated formatted Word doc: {out_path}')
