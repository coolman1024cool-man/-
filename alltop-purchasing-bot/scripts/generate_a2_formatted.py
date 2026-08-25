import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import os

template_path = r'D:\01邱正彥1130801-\01經費\06未結案\115年度-上半年\11150505-Usr-02-Hub7-吳仁明-自動化動作視覺控制系統軟體\A2規格表-自動化動作視覺控制系統軟體.docx'
output_dir = r'D:\01邱正彥1130801-\01經費\06未結案\115年度-下半年\子計畫2-陳隆泰-'

quotes = [
    {
        'filename': 'A2規格表-AI_Box主機套裝.docx',
        'title': '智能咖啡烘豆機AI Box主機套裝',
        'spec': '※智能咖啡烘豆機AI Box套裝\n1. AI Box主機組合(RK-reComputer工業版J4012套件 (安裝Jetson Orin NX 16 GB運算模組 + 128GB SSD))\n2. AI Box套件(鋁合金外殼適用於 Jetson Orin Nano 和 Jetson Orin NX 套件)'
    },
    {
        'filename': 'A2規格表-烘豆機MCU控制組材料.docx',
        'title': '烘豆機MCU控制組材料',
        'spec': '※智能咖啡烘豆機AI Box套裝\n1. MCU控制組(ESP32)\n2. 觸碰螢幕套件(WaveShare 樹莓派 7吋電容式觸碰螢幕)與HDMI線組\n3. 鏡頭模組套件(圓剛 PW310P 1080p)'
    },
    {
        'filename': 'A2規格表-烘豆機周邊與驅動套件.docx',
        'title': '智能咖啡烘豆機周邊與驅動套件',
        'spec': '※智能咖啡烘豆機AI Box套裝\n1. 瓦斯閥馬達驅動組(drv 8825)\n2. 麥克風套件(ReSpeaker Mic Array v3.0 (4麥克風聲學陣列))與延伸桿套件組\n3. 110V AC轉12V DC變壓器\n4. 落地式支架與手機夾'
    }
]

def set_font(run, font_name='標楷體', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

for q in quotes:
    doc = docx.Document(template_path)
    
    # 修改邊界 (上下各 2 公分 = 20mm = 0.7874 inches = 567 twips)
    for section in doc.sections:
        section.top_margin = docx.shared.Cm(2)
        section.bottom_margin = docx.shared.Cm(2)

    # 替換段落中的名稱與日期
    for p in doc.paragraphs:
        if '採購名稱' in p.text:
            p.clear()
            run1 = p.add_run(f'採購名稱︰{q["title"]}')
            set_font(run1, '標楷體', 12)
            
            # 使用 tab 或空格對齊
            run2 = p.add_run('                     日期︰115年8月25日')
            set_font(run2, '標楷體', 12)
    
    # 替換表格內容
    table = doc.tables[0]
    cell = table.cell(1, 0)
    cell.text = ''
    run_spec = cell.paragraphs[0].add_run(q['spec'])
    set_font(run_spec, '標楷體', 12)
    
    cell_unit = table.cell(1, 1)
    cell_unit.text = ''
    run_unit = cell_unit.paragraphs[0].add_run('1式')
    set_font(run_unit, '標楷體', 12)
    
    # 儲存檔案
    doc.save(os.path.join(output_dir, q['filename']))

print('Done Formatting')
