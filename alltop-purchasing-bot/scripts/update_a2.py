import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import os

target_dir = r'D:\01邱正彥1130801-\01經費\06未結案\115年度-下半年\子計畫2-陳隆泰-'

# 1. 刪除舊版 (新細明體) 的檔案
old_files = [
    'A2規格表-AI_Box主機套件.docx',
    'A2規格表-AI_Box_MCU與周邊套件.docx',
    'A2規格表-AI_Box_驅動與支架套件.docx'
]
for f in old_files:
    path = os.path.join(target_dir, f)
    if os.path.exists(path):
        os.remove(path)
        print(f"Deleted: {f}")

# 2. 修改新版 (標楷體) 檔案的內容
new_files = [
    'A2規格表-AI_Box主機套裝.docx',
    'A2規格表-烘豆機MCU控制組材料.docx',
    'A2規格表-烘豆機周邊與驅動套件.docx'
]

def set_font(run, font_name='標楷體', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

for f in new_files:
    path = os.path.join(target_dir, f)
    if not os.path.exists(path):
        continue
    
    doc = docx.Document(path)
    
    for p in doc.paragraphs:
        if '陳芸香' in p.text:
            text = p.text.replace('陳芸香', '徐佳琪')
            p.clear()
            run = p.add_run(text)
            set_font(run)
            
        if '交貨地點' in p.text:
            # 替換地點為 USR中心 (保留原本的分機格式)
            text = '三、交貨地點︰USR中心；分機號碼： 2800/2801 。'
            p.clear()
            run = p.add_run(text)
            set_font(run)
            
    doc.save(path)
    print(f"Updated: {f}")

print('Done')
