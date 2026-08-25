import docx
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import os

source_template = r'D:\01邱正彥1130801-\01經費\06未結案\115年度-上半年\11150505-Usr-02-Hub7-吳仁明-自動化動作視覺控制系統軟體\A2規格表-自動化動作視覺控制系統軟體.docx'
dest_template = r'D:\01邱正彥1130801-\01經費\ALLTOP_採購自動化工具\A2規格表_Template.docx'

def set_font(run, font_name='標楷體', size=12):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

doc = docx.Document(source_template)

# 設定邊界
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# 更新內容為完美初始狀態
for p in doc.paragraphs:
    if '陳芸香' in p.text:
        text = p.text.replace('陳芸香', '徐佳琪')
        p.clear()
        run = p.add_run(text)
        set_font(run)
        
    if '交貨地點' in p.text:
        text = '三、交貨地點︰USR中心；分機號碼： 2800/2801 。'
        p.clear()
        run = p.add_run(text)
        set_font(run)
        
    # 清空採購名稱，用佔位符
    if '採購名稱' in p.text:
        p.clear()
        run = p.add_run('採購名稱︰{{TITLE}}                     日期︰{{DATE}}')
        set_font(run)

# 清空表格，用佔位符
table = doc.tables[0]
cell_spec = table.cell(1, 0)
cell_spec.text = ''
run_spec = cell_spec.paragraphs[0].add_run('{{SPEC}}')
set_font(run_spec)

cell_unit = table.cell(1, 1)
cell_unit.text = ''
run_unit = cell_unit.paragraphs[0].add_run('{{UNIT}}')
set_font(run_unit)

doc.save(dest_template)
print('Template created successfully!')
